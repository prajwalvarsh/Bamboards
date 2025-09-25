#!/usr/bin/env python3
"""
Content Extractor Module for Bamboards
Extracts text content from various file formats for wordcloud generation
"""

import re
from pathlib import Path
from typing import List, Dict, Optional
import json


class ContentExtractor:
    def __init__(self):
        # German and English stopwords (basic set)
        self.stop_words = {
            # German stopwords
            'der', 'die', 'das', 'und', 'oder', 'aber', 'auch', 'noch', 'nicht',
            'ist', 'sind', 'war', 'waren', 'haben', 'hat', 'hatte', 'hatten',
            'werden', 'wird', 'wurde', 'wurden', 'sein', 'seine', 'seiner',
            'ich', 'du', 'er', 'sie', 'es', 'wir', 'ihr', 'sie', 'mich', 'dich',
            'sich', 'uns', 'euch', 'ihm', 'ihr', 'ihnen', 'mir', 'dir',
            'ein', 'eine', 'einer', 'eines', 'einem', 'einen',
            'auf', 'aus', 'bei', 'mit', 'nach', 'von', 'zu', 'an', 'in', 'für',
            'über', 'unter', 'durch', 'gegen', 'ohne', 'um', 'vor', 'zwischen',
            'dass', 'wenn', 'weil', 'da', 'als', 'wie', 'wo', 'was', 'wer',
            'welche', 'welcher', 'welches', 'dieser', 'diese', 'dieses',
            'jeder', 'jede', 'jedes', 'alle', 'alles', 'viele', 'wenige',
            'mehr', 'weniger', 'sehr', 'ganz', 'gar', 'nur', 'schon',
            
            # English stopwords
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
            'before', 'after', 'above', 'below', 'between', 'among', 'throughout',
            'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
            'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
            'must', 'can', 'shall', 'ought', 'need', 'dare',
            'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us',
            'them', 'my', 'your', 'his', 'her', 'its', 'our', 'their', 'mine',
            'yours', 'ours', 'theirs', 'myself', 'yourself', 'himself', 'herself',
            'itself', 'ourselves', 'yourselves', 'themselves',
            'this', 'that', 'these', 'those', 'what', 'which', 'who', 'whom',
            'whose', 'where', 'when', 'why', 'how', 'all', 'any', 'both', 'each',
            'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not',
            'only', 'own', 'same', 'so', 'than', 'too', 'very', 'just', 'now'
        }
        
        # Additional domain-specific stopwords
        self.domain_stopwords = {
            'bamboard', 'bamboards', 'display', 'screen', 'digital', 'public',
            'system', 'user', 'users', 'interface', 'design', 'technology',
            'page', 'document', 'file', 'pdf', 'docx', 'text', 'content'
        }
        
        self.all_stopwords = self.stop_words.union(self.domain_stopwords)
    
    def extract_from_pdf(self, filepath: Path) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except ImportError:
            print("  PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f" Error extracting from PDF {filepath.name}: {e}")
            return ""
    
    def extract_from_docx(self, filepath: Path) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(filepath)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except ImportError:
            print("  python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            print(f" Error extracting from DOCX {filepath.name}: {e}")
            return ""
    
    def extract_from_doc(self, filepath: Path) -> str:
        """Extract text from DOC file (older Word format)"""
        try:
            import subprocess
            # Try using antiword if available
            result = subprocess.run(['antiword', str(filepath)], 
                                  capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout
            else:
                print(f"  Could not extract from DOC file {filepath.name}")
                print(" Consider converting to DOCX format manually")
                return ""
        except FileNotFoundError:
            print(f"  antiword not found. Cannot extract from DOC file {filepath.name}")
            print(" Install antiword or convert to DOCX format")
            return ""
        except Exception as e:
            print(f"❌ Error extracting from DOC {filepath.name}: {e}")
            return ""
    
    def extract_from_txt(self, filepath: Path) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(filepath, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"❌ Error extracting from TXT {filepath.name}: {e}")
                return ""
    
    def extract_text(self, filepath: Path) -> str:
        """Extract text based on file extension"""
        extension = filepath.suffix.lower()
        
        print(f"📄 Extracting text from: {filepath.name}")
        
        if extension == '.pdf':
            return self.extract_from_pdf(filepath)
        elif extension == '.docx':
            return self.extract_from_docx(filepath)
        elif extension == '.doc':
            return self.extract_from_doc(filepath)
        elif extension in ['.txt', '.rtf']:
            return self.extract_from_txt(filepath)
        else:
            print(f" Unsupported file type: {extension}")
            return ""
    
    def clean_text(self, text: str) -> str:
        """Clean and preprocess text"""
        if not text:
            return ""
        
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep German umlauts and basic punctuation
        text = re.sub(r'[^\w\säöüÄÖÜß.,!?;:\-]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Convert to lowercase for processing
        text = text.lower()
        
        return text.strip()
    
    def extract_keywords(self, text: str, min_length: int = 3, max_length: int = 20) -> List[str]:
        """Extract meaningful keywords from text"""
        cleaned_text = self.clean_text(text)
        
        if not cleaned_text:
            return []
        
        # Split into words
        words = re.findall(r'\b[a-zA-ZäöüÄÖÜß]+\b', cleaned_text)
        
        # Filter words
        keywords = []
        for word in words:
            word_lower = word.lower()
            if (min_length <= len(word) <= max_length and 
                word_lower not in self.all_stopwords and 
                not word.isdigit() and
                not re.match(r'^[a-z]{1,2}$', word_lower)):  # Skip very short words
                keywords.append(word_lower)
        
        return keywords
    
    def extract_sentences_with_keywords(self, text: str, keywords: List[str]) -> List[str]:
        """Extract sentences that contain specific keywords"""
        if not text or not keywords:
            return []
        
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        
        relevant_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 20:  # Skip very short sentences
                sentence_lower = sentence.lower()
                if any(keyword in sentence_lower for keyword in keywords):
                    relevant_sentences.append(sentence)
        
        return relevant_sentences
    
    def process_file(self, file_info: Dict) -> Dict:
        """Process a single file and extract content"""
        filepath = file_info['path']
        
        result = {
            'filename': file_info['name'],
            'filepath': str(filepath),
            'is_interview_related': file_info['is_interview_related'],
            'raw_text': '',
            'cleaned_text': '',
            'keywords': [],
            'word_count': 0,
            'keyword_count': 0,
            'extraction_success': False
        }
        
        try:
            # Extract raw text
            raw_text = self.extract_text(filepath)
            
            if raw_text:
                result['raw_text'] = raw_text
                result['cleaned_text'] = self.clean_text(raw_text)
                result['keywords'] = self.extract_keywords(raw_text)
                result['word_count'] = len(raw_text.split())
                result['keyword_count'] = len(result['keywords'])
                result['extraction_success'] = True
                
                print(f" Extracted {result['word_count']} words, {result['keyword_count']} keywords")
            else:
                print(f"  No text extracted")
                
        except Exception as e:
            print(f"  Error processing file: {e}")
        
        return result
    
    def process_files(self, categorized_files: Dict[str, List[Dict]]) -> Dict:
        """Process all files and extract content"""
        print("\n📝 Content Extraction - Step 2")
        print("=" * 50)
        
        results = {
            'interview_content': [],
            'other_content': [],
            'all_content': [],
            'summary': {
                'total_files_processed': 0,
                'successful_extractions': 0,
                'total_words': 0,
                'total_keywords': 0,
                'interview_files_count': 0,
                'other_files_count': 0
            }
        }
        
        # Process interview files only (research papers and UX docs are excluded)
        print("\n Processing interview-related files...")
        for file_info in categorized_files.get('interview_files', []):
            content = self.process_file(file_info)
            results['interview_content'].append(content)
            results['all_content'].append(content)
            
            if content['extraction_success']:
                results['summary']['successful_extractions'] += 1
                results['summary']['total_words'] += content['word_count']
                results['summary']['total_keywords'] += content['keyword_count']
                results['summary']['interview_files_count'] += 1
        
        results['summary']['total_files_processed'] = len(results['all_content'])
        
        return results
    
    def save_extracted_content(self, results: Dict, output_file: str = "extracted_content.json"):
        """Save extracted content to JSON file"""
        output_path = Path(output_file)
        
        # Create a simplified version for JSON serialization
        json_results = {
            'summary': results['summary'],
            'interview_files': [
                {
                    'filename': content['filename'],
                    'word_count': content['word_count'],
                    'keyword_count': content['keyword_count'],
                    'extraction_success': content['extraction_success']
                }
                for content in results['interview_content']
            ],
            'other_files': [
                {
                    'filename': content['filename'],
                    'word_count': content['word_count'],
                    'keyword_count': content['keyword_count'],
                    'extraction_success': content['extraction_success']
                }
                for content in results['other_content']
            ]
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_results, f, indent=2, ensure_ascii=False)
        
        print(f"💾 Saved extraction summary to: {output_path}")
    
    def print_extraction_summary(self, results: Dict):
        """Print summary of extraction results"""
        summary = results['summary']
        
        print(f"\n Content Extraction Summary:")
        print(f"    Total files processed: {summary['total_files_processed']}")
        print(f"    Successful extractions: {summary['successful_extractions']}")
        print(f"   Total words extracted: {summary['total_words']:,}")
        print(f"   Total keywords found: {summary['total_keywords']:,}")
        print(f"   Interview files: {summary['interview_files_count']}")
        print(f"    Other files: {summary['other_files_count']}")
        


def main():
    """Test content extraction with existing filtered files"""
    from file_filter import FileFilter
    
    # First run file filtering to get categorized files
    SHARE_URL = "https://cloud.smartcitybamberg.de/s/fna28j9bAedqzP2"
    filter_tool = FileFilter(SHARE_URL)
    
    # Check if we have existing extracted files
    zip_file = Path("downloads/bamboards_files.zip")
    if zip_file.exists():
        print(" Using existing downloaded files...")
        categorized = filter_tool.extract_and_filter_zip(zip_file)
    else:
        print(" No downloaded files found. Run file_filter.py first.")
        return
    
 
    extractor = ContentExtractor()
    results = extractor.process_files(categorized)
    # Save results
    extractor.save_extracted_content(results)
    
    
    
    return results


if __name__ == "__main__":
    main()